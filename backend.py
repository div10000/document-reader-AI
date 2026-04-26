import fitz  # PyMuPDF
import re # regular expressions for text cleaning
import unicodedata # for text normalization
import docx
import pandas as pd
import os
from rank_bm25 import BM25Okapi

def extract_pages(file_path):
    """Reads the document based on its file extension and extracts text."""
    ext = os.path.splitext(file_path)[1].lower()
    pages = []
    
    # --- PDF Handling ---
    if ext == ".pdf":
        doc = fitz.open(file_path)
        for page_num in range(len(doc)):
            text = doc.load_page(page_num).get_text("text")
            pages.append({"page_num": str(page_num + 1), "content": text})
            
    # --- Word Document Handling ---
    elif ext == ".docx":
        doc = docx.Document(file_path)
        all_text_blocks = []
        
        # 1. Grab all standard paragraphs
        for para in doc.paragraphs:
            if para.text.strip():
                all_text_blocks.append(para.text.strip())
                
        # 2. Grab all text hiding inside tables
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    if cell.text.strip():
                        all_text_blocks.append(cell.text.strip())
        
        # 3. Chunk it into sections
        section_counter = 1
        # Group every 20 blocks (paragraphs/cells) into a "Section"
        for i in range(0, len(all_text_blocks), 20):
            chunk = all_text_blocks[i:i+20]
            pages.append({
                "page_num": f"Section {section_counter}", 
                "content": "\n".join(chunk)
            })
            section_counter += 1
            
    # --- Excel Handling ---
    elif ext in [".xlsx", ".xls"]:
        # Read all sheets into a dictionary of DataFrames
        xls = pd.read_excel(file_path, sheet_name=None)
        for sheet_name, df in xls.items():
            # Convert the dataframe table into a readable string format
            text = df.to_string(index=False)
            pages.append({"page_num": f"Sheet: {sheet_name}", "content": text})
            
    # --- Plain Text Handling ---
    elif ext == ".txt":
        with open(file_path, "r", encoding="utf-8") as f:
            text = f.read()
            pages.append({"page_num": "1", "content": text})
            
    else:
        raise ValueError(f"Unsupported file type: {ext}")
        
    return pages

def tokenize(text):
    """Cleans punctuation and splits text into words."""
    # 1. Convert to lowercase
    text = text.lower()
    
    # 2. Remove accents and special marks ---
    
    # This separates the base letter from the accent mark, and then throws away all the extra marks.
    text = unicodedata.normalize('NFKD', text).encode('ASCII', 'ignore').decode('utf-8')
    
    #  Remove all punctuation (Keep only letters, numbers, and spaces)
    # The regex [^\w\s] means "anything that is NOT an alphanumeric character or space"
    text = re.sub(r'[^\w\s]', '', text)
    
    # 3. Split into individual words
    return text.split()

def build_bm25_index(pages):
    """Creates the BM25 search index from the pages."""
    tokenized_corpus = [tokenize(page["content"]) for page in pages]
    return BM25Okapi(tokenized_corpus)

def get_answer(user_question, bm25_index, pages, llm, chat_history):
    """Searches for relevant pages and asks Groq, while remembering the past 4 messages."""
    tokenized_query = tokenize(user_question)
    
    # Grab the top n most relevant pages based on the NEW question
    best_pages = bm25_index.get_top_n(tokenized_query, pages, n=8)
    
    # Build the document context string
    context = ""
    for page in best_pages:
        context += f"\n--- Page {page['page_num']} ---\n{page['content']}\n"
        
    # --- NEW: Build the Chat History String ---
    history_text = ""
    for msg in chat_history:
        role = "User" if msg["role"] == "user" else "Assistant"
        history_text += f"{role}: {msg['content']}\n"
    
    # --- NEW: Updated Prompt with Memory ---
    prompt = f"""You are a helpful assistant analyzing a document.
Based ONLY on the Document Context below, answer the user's question.
Always mention the page numbers of the document where you found the answer.
If you don't know, say you don't know. Do NOT make up answers.
You also have access to the Chat History to understand what the user is referring to (like "it", "they", or "that").
If a question is aksed from context of previous response, do not read things from the provided "Document Context". Instead, try to answer from provided "Chat History".

Chat History (Last few messages):
{history_text}

Document Context:
{context}

Current Question: {user_question}
"""
    response = llm.invoke(prompt)
    
    return response.content, [p['page_num'] for p in best_pages]